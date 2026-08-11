from __future__ import annotations

import csv
import base64
import os
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


@dataclass(frozen=True)
class TextBlock:
    text: str
    source_file: str
    location: str


def _clean(value: object) -> str:
    return " ".join(str(value or "").replace("\x00", " ").split())


def _excel_blocks(path: Path) -> list[TextBlock]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    blocks: list[TextBlock] = []
    try:
        for sheet in workbook.worksheets:
            headers: list[str] = []
            for row_number, row in enumerate(sheet.iter_rows(values_only=True), 1):
                cells = [_clean(cell) for cell in row]
                if not headers and any(name in cell for cell in cells for name in ("项目名称", "包号", "包名称", "采购范围", "分标编号")):
                    headers = cells
                    text = " | ".join(filter(None, cells))
                elif headers:
                    text = " | ".join(
                        f"{headers[index] or f'字段{index + 1}'}：{value}"
                        for index, value in enumerate(cells)
                        if value
                    )
                else:
                    text = " | ".join(filter(None, cells))
                if text:
                    blocks.append(TextBlock(text, path.name, f"工作表“{sheet.title}”第 {row_number} 行"))
    finally:
        workbook.close()
    return blocks


def _word_blocks(path: Path) -> list[TextBlock]:
    document = Document(path)
    blocks = [TextBlock(_clean(p.text), path.name, f"段落 {index}") for index, p in enumerate(document.paragraphs, 1) if _clean(p.text)]
    for table_number, table in enumerate(document.tables, 1):
        for row_number, row in enumerate(table.rows, 1):
            text = " | ".join(_clean(cell.text) for cell in row.cells if _clean(cell.text))
            if text:
                blocks.append(TextBlock(text, path.name, f"表格 {table_number} 第 {row_number} 行"))
    return blocks


def _ocr_pdf_page(page, source_name: str, page_number: int) -> TextBlock | None:
    executable = shutil.which("tesseract")
    if not executable:
        return None
    with tempfile.TemporaryDirectory(prefix="sgcc-ocr-") as temporary:
        image_path = Path(temporary) / f"page-{page_number}.png"
        page.get_pixmap(matrix=__import__("fitz").Matrix(2, 2), alpha=False).save(image_path)
        completed = subprocess.run(
            [executable, str(image_path), "stdout", "-l", os.getenv("OCR_LANGUAGES", "chi_sim+eng"), "--psm", "6"],
            capture_output=True, timeout=45, check=False,
        )
        text = _clean(completed.stdout.decode("utf-8", errors="replace"))
        return TextBlock(text, source_name, f"第 {page_number} 页（OCR）") if text else None


def _pdf_blocks(path: Path) -> tuple[list[TextBlock], str]:
    import fitz

    blocks: list[TextBlock] = []
    warnings: list[str] = []
    ocr_limit = max(0, min(int(os.getenv("OCR_MAX_PAGES", "30")), 100))
    with fitz.open(path) as document:
        for page_number, page in enumerate(document, 1):
            text = _clean(page.get_text("text"))
            if len(text) >= 30:
                blocks.append(TextBlock(text, path.name, f"第 {page_number} 页"))
                continue
            if page_number > ocr_limit:
                warnings.append(f"{path.name} 超过 OCR 页数上限 {ocr_limit}，后续扫描页未识别")
                continue
            try:
                ocr_block = _ocr_pdf_page(page, path.name, page_number)
                if ocr_block:
                    blocks.append(ocr_block)
                elif not shutil.which("tesseract"):
                    warnings.append("未安装 Tesseract，扫描 PDF 无法 OCR")
                else:
                    warnings.append(f"{path.name} 第 {page_number} 页 OCR 未提取到文字")
            except subprocess.TimeoutExpired:
                warnings.append(f"{path.name} 第 {page_number} 页 OCR 超时")
            except Exception as exc:
                warnings.append(f"{path.name} 第 {page_number} 页 OCR 失败：{type(exc).__name__}")
    return blocks, "；".join(dict.fromkeys(warnings))


def _text_blocks(path: Path) -> list[TextBlock]:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        return []
    return [TextBlock(_clean(line), path.name, f"第 {index} 行") for index, line in enumerate(text.splitlines(), 1) if _clean(line)]


def _converted_blocks(original: Path, converted: Path, label: str) -> tuple[list[TextBlock], str]:
    blocks, warning = parse_document(converted)
    rewritten = [TextBlock(block.text, original.name, f"{block.location}（由 {label} 转换）") for block in blocks]
    return rewritten, warning


def _convert_legacy_office(path: Path) -> tuple[list[TextBlock], str]:
    executable = shutil.which("libreoffice") or shutil.which("soffice")
    if not executable:
        return [], f"未安装 LibreOffice，无法转换 {path.suffix.casefold()} 文件"
    target_format = "xlsx" if path.suffix.casefold() == ".xls" else "docx"
    with tempfile.TemporaryDirectory(prefix="sgcc-office-") as temporary:
        output_dir = Path(temporary) / "output"
        profile_dir = Path(temporary) / "profile"
        output_dir.mkdir()
        completed = subprocess.run(
            [
                executable,
                "--headless",
                f"-env:UserInstallation={profile_dir.resolve().as_uri()}",
                "--convert-to",
                target_format,
                "--outdir",
                str(output_dir),
                str(path),
            ],
            capture_output=True,
            timeout=120,
            check=False,
        )
        converted = output_dir / f"{path.stem}.{target_format}"
        if completed.returncode != 0 or not converted.exists():
            return [], f"{path.name} 经 LibreOffice 转换失败"
        return _converted_blocks(path, converted, f"{path.suffix.upper()}→{target_format.upper()}")


def _convert_ofd(path: Path) -> tuple[list[TextBlock], str]:
    try:
        from easyofd.ofd import OFD
    except ImportError:
        return [], "未安装 easyofd，无法将 OFD 转换为 PDF"
    with tempfile.TemporaryDirectory(prefix="sgcc-ofd-") as temporary:
        converted = Path(temporary) / f"{path.stem}.pdf"
        converter = OFD()
        try:
            converter.read(base64.b64encode(path.read_bytes()).decode("ascii"), save_xml=False)
            converted.write_bytes(converter.to_pdf())
        except Exception as exc:
            return [], f"{path.name} 转换 PDF 失败：{type(exc).__name__}"
        finally:
            try:
                converter.del_data()
            except Exception:
                pass
        return _converted_blocks(path, converted, "OFD→PDF")


def parse_document(path: Path) -> tuple[list[TextBlock], str]:
    suffix = path.suffix.casefold()
    try:
        if suffix == ".xlsx":
            return _excel_blocks(path), ""
        if suffix == ".docx":
            return _word_blocks(path), ""
        if suffix == ".pdf":
            blocks, warning = _pdf_blocks(path)
            return blocks, warning if blocks else (warning or "PDF 未提取到可读文本")
        if suffix in {".txt", ".csv"}:
            return _text_blocks(path), ""
        if suffix in {".xls", ".doc"}:
            return _convert_legacy_office(path)
        if suffix == ".ofd":
            return _convert_ofd(path)
        return [], ""
    except Exception as exc:
        return [], f"{path.name} 解析失败：{type(exc).__name__}"
