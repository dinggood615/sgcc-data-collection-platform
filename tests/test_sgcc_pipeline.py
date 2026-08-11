from io import BytesIO

from openpyxl import Workbook
from docx import Document

from app.database import connect, init_db
from app.sgcc.pipeline import ingest_attachment
from app.sgcc import pipeline


def _workbook_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "分包清单"
    sheet.append(["分标编号", "包号", "项目名称", "采购范围"])
    sheet.append(["SGCC-2026-01", "包1", "信息系统升级项目", "应用软件开发及实施服务"])
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def test_xlsx_is_matched_and_keeps_row_evidence(monkeypatch, tmp_path):
    database = tmp_path / "platform.sqlite3"
    monkeypatch.setenv("DATABASE_PATH", str(database))
    init_db()
    result = ingest_attachment(
        "采购清单.xlsx",
        _workbook_bytes(),
        "NOTICE-1",
        "https://ecp.sgcc.com.cn/example",
        ["软件开发", "信息化"],
        [],
    )
    assert result.status == "processed"
    assert result.matched_count >= 1
    with connect() as db:
        row = db.execute("SELECT * FROM sgcc_packages WHERE relevance_score>=20 ORDER BY relevance_score DESC LIMIT 1").fetchone()
    assert row is not None
    assert "第 2 行" in row["source_location"]
    assert "应用软件开发" in row["evidence"]


def test_identical_attachment_is_not_processed_twice(monkeypatch, tmp_path):
    monkeypatch.setenv("DATABASE_PATH", str(tmp_path / "platform.sqlite3"))
    init_db()
    payload = _workbook_bytes()
    first = ingest_attachment("采购清单.xlsx", payload, "NOTICE-1", "", ["软件开发"], [])
    second = ingest_attachment("采购清单.xlsx", payload, "NOTICE-1", "", ["软件开发"], [])
    assert first.status == "processed"
    assert second.status == "duplicate"


def test_word_fields_on_adjacent_paragraphs_are_joined(monkeypatch, tmp_path):
    monkeypatch.setenv("DATABASE_PATH", str(tmp_path / "platform.sqlite3"))
    init_db()
    document = Document()
    document.add_paragraph("分标编号：SGCC-2026-02")
    document.add_paragraph("包号：包3")
    document.add_paragraph("项目名称：统一数据中台建设")
    document.add_paragraph("服务内容：应用软件开发、数据治理与实施服务")
    output = BytesIO()
    document.save(output)
    result = ingest_attachment("技术规范.docx", output.getvalue(), "NOTICE-2", "", ["软件开发", "数据治理"], [])
    assert result.matched_count == 1
    with connect() as db:
        row = db.execute("SELECT * FROM sgcc_packages WHERE relevance_score>=20").fetchone()
    assert row["package_no"] == "包3"
    assert row["project_name"] == "统一数据中台建设"
    assert "软件开发" in row["evidence"]


def test_attachment_documents_are_parsed_in_bounded_parallel_pool(monkeypatch, tmp_path):
    paths = []
    for index in range(4):
        path = tmp_path / f"{index}.txt"
        path.write_text(f"项目 {index}", encoding="utf-8")
        paths.append(path)
    seen = []

    def fake_parse(path):
        seen.append(path.name)
        return [pipeline.TextBlock(path.stem, path.name, "测试")], ""

    monkeypatch.setenv("ATTACHMENT_PARSE_WORKERS", "4")
    monkeypatch.setattr(pipeline, "parse_document", fake_parse)
    results = pipeline._parse_paths_parallel(paths)
    assert len(results) == 4
    assert set(seen) == {"0.txt", "1.txt", "2.txt", "3.txt"}


def test_unreadable_attachment_is_automatically_classified(monkeypatch, tmp_path):
    monkeypatch.setenv("DATABASE_PATH", str(tmp_path / "platform.sqlite3"))
    init_db()
    result = ingest_attachment("未知格式.bin", b"not-readable", "NOTICE-3", "", ["软件"], [])
    assert result.status == "no_text"
    with connect() as db:
        row = db.execute("SELECT status,message FROM sgcc_documents WHERE sha256=?", (result.sha256,)).fetchone()
    assert row["status"] == "no_text"
    assert "自动处理完成" in row["message"]
